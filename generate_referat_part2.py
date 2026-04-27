"""
Generate referat (course work) Part 2 – Server/Backend side
of the "Matrix of Competencies" project.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ──────────────────────────────────────────────
# Helper utilities
# ──────────────────────────────────────────────

def set_para_format(para, first_line_indent=True, space_before=0,
                    space_after=6, line_spacing=1.5):
    """Apply standard paragraph formatting (ГОСТ-like)."""
    pf = para.paragraph_format
    if first_line_indent:
        pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def add_paragraph(doc, text, style='Normal', bold=False, indent=True,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=14):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    set_para_format(p, first_line_indent=indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def heading(doc, text, level=1):
    """Add a numbered heading that looks like chapter/section titles."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p


def sub_heading(doc, text):
    return heading(doc, text, level=2)


def add_page_break(doc):
    doc.add_page_break()


def add_table_of_contents_entry(doc, text, dots=True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    return p


# ──────────────────────────────────────────────
# Document initialisation
# ──────────────────────────────────────────────

doc = Document()

# Page margins (ГОСТ 7.32)
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)

# Default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

# ══════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════

def centered(doc, text, bold=False, size=14, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return p

centered(doc, 'МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ', bold=True, size=12)
centered(doc, 'Федеральное государственное автономное образовательное учреждение', size=12)
centered(doc, 'высшего образования', size=12)
centered(doc, '«НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ УНИВЕРСИТЕТ»', bold=True, size=12)
centered(doc, '', size=12)
centered(doc, 'Кафедра информационных технологий', size=12)
centered(doc, '', size=12, space_before=60)

centered(doc, 'КУРСОВАЯ РАБОТА', bold=True, size=16, space_before=40)
centered(doc, '', size=12)
centered(doc, 'по дисциплине «Проектирование информационных систем»', size=14)
centered(doc, '', size=12, space_before=10)
centered(doc,
         'Тема: Проектирование и разработка серверной части системы управления\nкомпетенциями и задачами сотрудников «Матрица компетенций»',
         bold=True, size=14, space_before=10, space_after=40)

# Signature block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.first_line_indent = Cm(0)
p.paragraph_format.space_before = Pt(60)
r = p.add_run('Выполнил: студент группы ИТ-41\nИванов И.И.\n\nПроверил: к.т.н., доцент\nПетров П.П.')
r.font.size = Pt(14)
r.font.name = 'Times New Roman'

centered(doc, 'Москва, 2024', size=14, space_before=80)

add_page_break(doc)

# ══════════════════════════════════════════════
# РЕФЕРАТ
# ══════════════════════════════════════════════

heading(doc, 'РЕФЕРАТ')

add_paragraph(doc,
    'Расчётно-пояснительная записка: 52 с., 11 рисунков, 3 таблицы, 15 источников.')

add_paragraph(doc,
    'Ключевые слова: FastAPI, Python, PostgreSQL, REST API, JWT, Docker, '
    'SQLAlchemy, Redis, Apache Airflow, dbt, серверная часть, '
    'микросервисная архитектура, компетенции, управление задачами.')

add_paragraph(doc,
    'Объектом разработки является серверная часть системы управления задачами '
    'и компетенциями сотрудников «Матрица компетенций», включающая REST API '
    'на базе FastAPI, реляционную базу данных PostgreSQL, кэш Redis, '
    'систему ETL-оркестрации Apache Airflow и BI-контур на основе Apache Superset.')

add_paragraph(doc,
    'Цель работы — спроектировать и реализовать масштабируемую серверную '
    'инфраструктуру, обеспечивающую хранение данных, бизнес-логику, '
    'аутентификацию и аналитический контур платформы «Матрица компетенций».')

add_paragraph(doc,
    'Поставленные задачи решаются путём проектирования и разработки '
    'слоёв REST API, базы данных, ETL-пайплайнов и контейнеризованной '
    'инфраструктуры на базе Docker Compose.')

add_page_break(doc)

# ══════════════════════════════════════════════
# СОДЕРЖАНИЕ
# ══════════════════════════════════════════════

heading(doc, 'СОДЕРЖАНИЕ')

toc_items = [
    ('ВВЕДЕНИЕ', '5'),
    ('1. МЕТОДЫ И ИНСТРУМЕНТЫ ПРОГРАММНОЙ ИНЖЕНЕРИИ', '8'),
    ('   1.1. Основные требования к серверной части', '8'),
    ('   1.2. Анализ аналогов и прототипов серверных решений', '10'),
    ('   1.3. Обоснование выбора инструментов для серверной части', '13'),
    ('   1.4. Обоснование выбора платформы для хранения данных', '16'),
    ('2. ПРОЕКТИРОВАНИЕ КОМПОНЕНТОВ ПРОГРАММНОГО ПРОДУКТА', '19'),
    ('   2.1. Проектирование функциональной модели', '19'),
    ('   2.2. Принцип работы приложения', '22'),
    ('   2.3. Структура базы данных', '26'),
    ('   2.4. Схемы основных алгоритмов', '30'),
    ('3. ТЕСТИРОВАНИЕ И ИНТЕГРАЦИЯ КОМПОНЕНТОВ', '34'),
    ('   3.1. Разработка тестов и результаты тестирования', '34'),
    ('   3.2. Интерфейс REST API (Swagger/OpenAPI)', '38'),
    ('   3.3. Интерфейс BI-контура (Apache Superset)', '40'),
    ('   3.4. Руководство пользователя API', '42'),
    ('   3.5. Руководство администратора', '45'),
    ('ЗАКЛЮЧЕНИЕ', '49'),
    ('СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', '50'),
]

for label, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(3)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT)
    r1 = p.add_run(label)
    r1.font.size = Pt(13)
    r1.font.name = 'Times New Roman'
    r2 = p.add_run('\t' + page)
    r2.font.size = Pt(13)
    r2.font.name = 'Times New Roman'

add_page_break(doc)

# ══════════════════════════════════════════════
# ВВЕДЕНИЕ
# ══════════════════════════════════════════════

heading(doc, 'ВВЕДЕНИЕ')

add_paragraph(doc,
    'Актуальность темы курсовой работы обусловлена стремительным ростом '
    'цифровизации корпоративных процессов и повышением требований к системам '
    'управления человеческим капиталом. Современные организации сталкиваются '
    'с необходимостью автоматизированного сопоставления компетенций сотрудников '
    'с требованиями задач, оперативного распределения нагрузки и '
    'аналитического мониторинга кадрового потенциала. Традиционные '
    'таблично-ориентированные инструменты не обеспечивают требуемой '
    'оперативности и интеграции с мобильными клиентами.')

add_paragraph(doc,
    'Проект «Матрица компетенций» представляет собой комплексную цифровую '
    'платформу, объединяющую мобильный клиент на React Native/Expo, '
    'REST API на FastAPI, реляционное хранилище данных PostgreSQL, '
    'кэш Redis, ETL-оркестрацию на Apache Airflow с dbt-трансформациями '
    'и BI-визуализацию через Apache Superset. Данная курсовая работа '
    'посвящена проектированию и реализации серверной части этой платформы.')

add_paragraph(doc,
    'Объектом курсовой работы является серверная инфраструктура системы '
    '«Матрица компетенций», включающая API-слой, слой данных, '
    'интеграционный ETL-контур и аналитический BI-контур.')

add_paragraph(doc,
    'Предметом исследования являются методы проектирования асинхронных '
    'REST API, подходы к организации реляционных схем баз данных '
    'для предметной области управления компетенциями, а также '
    'инструменты контейнеризации и оркестрации ETL-процессов.')

add_paragraph(doc,
    'Целью работы является проектирование и реализация масштабируемой, '
    'безопасной и легко сопровождаемой серверной части платформы '
    '«Матрица компетенций», обеспечивающей корректное взаимодействие '
    'с мобильным клиентом и аналитическим контуром.')

add_paragraph(doc, 'Для достижения поставленной цели решаются следующие задачи:')

tasks = [
    'провести анализ существующих решений и обосновать выбор технологического стека серверной части;',
    'спроектировать функциональную модель серверной части с использованием методологии UML;',
    'разработать схему базы данных PostgreSQL, обеспечивающую хранение данных пользователей, задач, компетенций и уведомлений;',
    'реализовать REST API на базе FastAPI с JWT-аутентификацией, ролевой моделью и слоёной архитектурой;',
    'спроектировать и реализовать ETL-пайплайн с использованием Apache Airflow и dbt;',
    'обеспечить контейнеризацию всех сервисов посредством Docker Compose;',
    'провести тестирование серверных компонентов и зафиксировать результаты.',
]
for i, t in enumerate(tasks, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t)
    r.font.size = Pt(14)
    r.font.name = 'Times New Roman'

add_paragraph(doc,
    'Практическая значимость работы заключается в том, что разработанная '
    'серверная часть интегрирована с мобильным клиентом и аналитическим '
    'контуром, а все компоненты развёртываются автоматически посредством '
    'Docker Compose, что существенно снижает операционные расходы.')

add_paragraph(doc,
    'Структура работы включает введение, три главы, заключение и список '
    'использованных источников. В первой главе рассмотрены методы и '
    'инструменты программной инженерии. Во второй — проектирование '
    'компонентов. В третьей — тестирование и интеграция. '
    'Общий объём работы составляет 52 страницы.')

add_page_break(doc)

# ══════════════════════════════════════════════
# ГЛАВА 1
# ══════════════════════════════════════════════

heading(doc, '1. МЕТОДЫ И ИНСТРУМЕНТЫ ПРОГРАММНОЙ ИНЖЕНЕРИИ')

sub_heading(doc, '1.1. Основные требования к серверной части')

add_paragraph(doc,
    'При проектировании серверной части системы «Матрица компетенций» '
    'были сформулированы функциональные и нефункциональные требования '
    'на основе анализа предметной области и пожеланий заказчика.')

add_paragraph(doc,
    'К функциональным требованиям относятся:')

func_reqs = [
    'регистрация и аутентификация пользователей с ролями «менеджер» и «сотрудник»;',
    'CRUD-операции над задачами (создание, чтение, обновление, удаление);',
    'управление профилями сотрудников и их компетенциями;',
    'подача и обработка заявок сотрудников на выполнение задач;',
    'формирование персонализированных рекомендаций по задачам на основе компетенций;',
    'отправка и получение уведомлений внутри системы;',
    'импорт исходных данных из Excel-файлов через ETL-пайплайн;',
    'предоставление аналитических витрин данных для BI-слоя;',
    'интеграция с мобильным клиентом по протоколу HTTP/JSON.',
]
for r_text in func_reqs:
    p = doc.add_paragraph('— ' + r_text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Нефункциональные требования определяют качество реализации и условия '
    'эксплуатации системы. К ним относятся следующие:')

nonfunc_reqs = [
    'Производительность: API должен обрабатывать не менее 100 одновременных запросов без деградации времени ответа свыше 500 мс.',
    'Безопасность: хранение паролей в хешированном виде (bcrypt), использование JWT-токенов с ограниченным сроком действия.',
    'Надёжность: автоматическая инициализация базы данных при первом запуске, fallback-механизмы при сбоях.',
    'Сопровождаемость: слоёная архитектура (router → service → model/repository), разделение ответственности.',
    'Переносимость: полная контейнеризация через Docker Compose.',
    'Масштабируемость: горизонтальное масштабирование API-инстансов за балансировщиком нагрузки.',
]
for r_text in nonfunc_reqs:
    p = doc.add_paragraph('— ' + r_text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Для структурирования требований использовался метод MoSCoW '
    '(Must Have / Should Have / Could Have / Wont Have). Критическими '
    '(Must Have) были признаны аутентификация, управление задачами и '
    'хранение компетенций. К категории Should Have отнесены уведомления '
    'и рекомендательная система. Could Have — расширенная аналитика '
    'и экспорт отчётов.')

sub_heading(doc, '1.2. Анализ аналогов и прототипов серверных решений')

add_paragraph(doc,
    'Перед началом проектирования был проведён анализ существующих систем '
    'управления задачами и компетенциями, а также технологических решений '
    'их серверных частей. Рассматривались как коммерческие, '
    'так и открытые платформы.')

add_paragraph(doc,
    'Jira (Atlassian) — широко распространённая система управления проектами '
    'и задачами. Серверная часть реализована на Java (Spring Boot), '
    'использует PostgreSQL/MySQL для хранения данных. Система обладает '
    'развитыми возможностями настройки рабочих процессов, REST API '
    'и системой плагинов. Недостатком является высокая стоимость '
    'лицензирования и избыточная сложность для задач управления '
    'компетенциями отдельного подразделения.')

add_paragraph(doc,
    'YouTrack (JetBrains) — система управления проектами с функциями '
    'тайм-трекинга. Реализована на Kotlin/JVM, поддерживает PostgreSQL '
    'и собственную облачную СУБД. Предоставляет REST API и Webhooks. '
    'Ограничение — отсутствие встроенного модуля компетенций '
    'и рекомендательной системы.')

add_paragraph(doc,
    'Redmine — открытая система управления проектами на Ruby on Rails. '
    'Поддерживает PostgreSQL, MySQL, SQLite. Предоставляет REST API '
    'с JSON-форматом. Плагины расширяют базовую функциональность, '
    'однако модуль компетенций требует значительной кастомизации.')

add_paragraph(doc,
    'Анализ показывает, что готовые решения либо избыточны по функциональности '
    'и стоимости, либо не содержат специализированного модуля управления '
    'компетенциями с рекомендательной логикой. Это обосновывает '
    'целесообразность разработки собственной серверной части.')

# Table 1 – comparison
heading(doc, 'Таблица 1 — Сравнительный анализ аналогов', level=2)

table1 = doc.add_table(rows=5, cols=5)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Критерий', 'Jira', 'YouTrack', 'Redmine', 'Матрица компетенций']
for i, h in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(12)
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

rows_data = [
    ['Управление задачами', '+', '+', '+', '+'],
    ['Модуль компетенций', '—', '—', 'частично', '+'],
    ['Рекомендательная система', '—', '—', '—', '+'],
    ['Открытый исходный код', '—', '—', '+', '+'],
]
for i, row_data in enumerate(rows_data, 1):
    for j, cell_text in enumerate(row_data):
        cell = table1.rows[i].cells[j]
        cell.text = cell_text
        cell.paragraphs[0].runs[0].font.size = Pt(12)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

doc.add_paragraph()  # spacing after table

sub_heading(doc, '1.3. Обоснование выбора инструментов для серверной части')

add_paragraph(doc,
    'На основании сформированных требований и анализа аналогов был обоснован '
    'следующий технологический стек серверной части.')

add_paragraph(doc,
    'Python 3.11+ выбран в качестве основного языка программирования '
    'благодаря богатой экосистеме библиотек для работы с данными, '
    'развитым фреймворкам для API и высокой скорости прототипирования. '
    'Python занимает первое место в индексе TIOBE (2023–2024) и '
    'является де-факто стандартом для задач Data Engineering и ML.')

add_paragraph(doc,
    'FastAPI — современный асинхронный веб-фреймворк для Python, '
    'основанный на стандарте ASGI и библиотеке Starlette. '
    'Ключевые преимущества: автоматическая генерация OpenAPI/Swagger '
    'документации, встроенная валидация данных через Pydantic, '
    'поддержка async/await, что критично для высоконагруженных API. '
    'По результатам независимых бенчмарков TechEmpower FastAPI '
    'демонстрирует производительность, сопоставимую с Node.js Express '
    'и превосходящую синхронный Django REST Framework.')

add_paragraph(doc,
    'SQLAlchemy 2.x (async) — наиболее зрелая ORM для Python '
    'с поддержкой асинхронного режима работы. Обеспечивает '
    'декларативное определение моделей, миграции через Alembic, '
    'пул соединений и type-safe запросы. Использование asyncpg-драйвера '
    'при работе с PostgreSQL позволяет достигать высокой пропускной '
    'способности при обработке запросов.')

add_paragraph(doc,
    'Pydantic v2 применяется для валидации входных и выходных данных '
    'API. Схемы Pydantic одновременно служат документацией контракта '
    'API и инструментом защиты от некорректных данных. '
    'Версия 2 написана на Rust, что обеспечивает многократное '
    'ускорение валидации по сравнению с версией 1.')

add_paragraph(doc,
    'JWT (JSON Web Tokens) используется для stateless-аутентификации '
    'и авторизации. Токены подписываются секретным ключом '
    '(HS256), содержат информацию о роли пользователя и '
    'имеют ограниченный срок действия. Это позволяет API-серверам '
    'работать без сохранения сессий, что упрощает горизонтальное '
    'масштабирование.')

add_paragraph(doc,
    'Redis 7 применяется как высокопроизводительное in-memory хранилище '
    'для кэширования часто запрашиваемых данных (профили пользователей, '
    'списки задач), хранения временных токенов и служебных очередей. '
    'Redis обеспечивает время доступа к данным менее 1 мс, '
    'что существенно снижает нагрузку на PostgreSQL.')

add_paragraph(doc,
    'Flask (flask_auth) используется как легковесный веб-модуль '
    'для обработки сценариев браузерной авторизации и перехода '
    'в панель Superset. Отделение этого модуля от основного API '
    'обеспечивает разделение ответственности.')

sub_heading(doc, '1.4. Обоснование выбора платформы для хранения и обработки данных')

add_paragraph(doc,
    'PostgreSQL 15 выбран в качестве основной СУБД по следующим причинам: '
    'поддержка JSONB-типов для гибкого хранения полуструктурированных данных, '
    'развитый механизм индексирования (B-tree, GIN, GiST), '
    'транзакционность ACID, поддержка оконных функций и CTE '
    'для аналитических запросов, зрелая экосистема инструментов '
    '(pgAdmin, pg_dump, Alembic). PostgreSQL является СУБД с открытым '
    'исходным кодом с наибольшим рейтингом среди реляционных БД '
    'по версии DB-Engines Ranking (2024).')

add_paragraph(doc,
    'Apache Airflow 2.x выбран как стандарт де-факто для оркестрации '
    'ETL/ELT пайплайнов в Python-экосистеме. DAG (Directed Acyclic Graph) '
    'позволяет декларативно описывать зависимости между задачами, '
    'планировать их выполнение по расписанию и отслеживать статус '
    'через встроенный веб-интерфейс. Airflow поддерживает интеграцию '
    'с PostgreSQL в качестве метаданных-хранилища.')

add_paragraph(doc,
    'dbt (data build tool) применяется для трансформации данных '
    'в аналитическом слое. dbt позволяет описывать SQL-трансформации '
    'в виде версионируемых моделей, автоматически строить граф '
    'зависимостей, тестировать данные и генерировать документацию. '
    'Это обеспечивает чистоту аналитического слоя и воспроизводимость '
    'результатов.')

add_paragraph(doc,
    'Apache Superset используется для BI-визуализации. '
    'Superset поддерживает прямое подключение к PostgreSQL, '
    'построение дашбордов без программирования, управление доступом '
    'на уровне ролей и экспорт данных. Он является открытой '
    'альтернативой Tableau и Power BI, что соответствует '
    'бюджетным ограничениям проекта.')

add_paragraph(doc,
    'Docker и Docker Compose обеспечивают воспроизводимое '
    'развёртывание всех сервисов. Каждый компонент платформы '
    'изолирован в отдельном контейнере с чётко определёнными '
    'зависимостями и переменными окружения. Единая bridge-сеть '
    'my-network обеспечивает межсервисное взаимодействие.')

add_page_break(doc)

# ══════════════════════════════════════════════
# ГЛАВА 2
# ══════════════════════════════════════════════

heading(doc, '2. ПРОЕКТИРОВАНИЕ КОМПОНЕНТОВ ПРОГРАММНОГО ПРОДУКТА')

sub_heading(doc, '2.1. Проектирование функциональной модели')

add_paragraph(doc,
    'Для формального описания функциональных требований к серверной '
    'части использован язык моделирования UML. Разработаны диаграммы '
    'прецедентов (Use Case), компонентов и последовательностей.')

add_paragraph(doc,
    'Диаграмма прецедентов описывает взаимодействие двух основных '
    'акторов — «Сотрудник» и «Менеджер» — с сервером. '
    'Дополнительными акторами являются «Система» (автоматические '
    'процессы Airflow) и «Администратор».')

add_paragraph(doc,
    'Основные прецеденты для актора «Сотрудник»:')
emp_cases = [
    'UC-01: Авторизация в системе',
    'UC-02: Просмотр списка доступных задач',
    'UC-03: Подача заявки на выполнение задачи',
    'UC-04: Просмотр персональных рекомендаций',
    'UC-05: Обновление профиля и компетенций',
    'UC-06: Просмотр уведомлений',
    'UC-07: Отслеживание статуса своих заявок',
]
for uc in emp_cases:
    p = doc.add_paragraph('— ' + uc)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Основные прецеденты для актора «Менеджер»:')
mgr_cases = [
    'UC-08: Создание новой задачи',
    'UC-09: Редактирование и удаление задачи',
    'UC-10: Назначение задачи сотруднику',
    'UC-11: Просмотр матрицы компетенций команды',
    'UC-12: Управление статусами задач',
    'UC-13: Просмотр BI-дашбордов',
    'UC-14: Управление уведомлениями',
]
for uc in mgr_cases:
    p = doc.add_paragraph('— ' + uc)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Диаграмма компонентов серверной части отражает следующую структуру: '
    'внешний клиент (mobile-app) взаимодействует с FastAPI через HTTPS; '
    'FastAPI включает компоненты Router (маршрутизация), Service '
    '(бизнес-логика) и Repository (работа с БД); Repository обращается '
    'к PostgreSQL через asyncpg; Redis используется Service-слоем '
    'для кэширования; Airflow + dbt взаимодействуют напрямую '
    'с PostgreSQL для ETL и трансформаций.')

add_paragraph(doc,
    'Диаграмма последовательностей для прецедента UC-01 (авторизация): '
    '1) Клиент отправляет POST /auth/login с логином и паролем; '
    '2) Router передаёт запрос в AuthService; '
    '3) AuthService запрашивает UserRepository; '
    '4) UserRepository выполняет SELECT к PostgreSQL; '
    '5) При совпадении хеша AuthService генерирует JWT; '
    '6) Router возвращает токен клиенту.')

sub_heading(doc, '2.2. Принцип работы приложения')

add_paragraph(doc,
    'Серверная часть «Матрицы компетенций» построена по принципу '
    'трёхуровневой архитектуры с дополнительным ETL- и BI-контурами. '
    'Рассмотрим каждый уровень подробнее.')

add_paragraph(doc,
    'Уровень маршрутизации (Routers). FastAPI-приложение содержит '
    'следующие маршрутизаторы:')
routers = [
    'auth_router (/auth) — регистрация, вход, обновление токена;',
    'employees_router (/employees) — управление профилями сотрудников;',
    'tasks_router (/tasks) — CRUD задач и управление статусами;',
    'notifications_router (/notifications) — получение и отметка уведомлений;',
    'recommendations_router (/recommendations) — персонализированные рекомендации.',
]
for r_text in routers:
    p = doc.add_paragraph('— ' + r_text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Уровень сервисов (Services) реализует бизнес-логику. '
    'Сервис рекомендаций (RecommendationService) анализирует '
    'компетенции сотрудника и открытые задачи, формируя ранжированный '
    'список подходящих задач на основе пересечения множеств навыков. '
    'TaskService управляет жизненным циклом задачи через конечный '
    'автомат состояний: pending → assigned → in_progress → '
    'review → completed/rejected.')

add_paragraph(doc,
    'Уровень репозиториев (Repositories) абстрагирует операции '
    'с базой данных. Каждая сущность (User, Task, Competency, '
    'Notification) имеет соответствующий репозиторий. '
    'Репозитории работают через SQLAlchemy AsyncSession, '
    'что обеспечивает неблокирующие операции с БД.')

add_paragraph(doc,
    'Старт приложения (main.py) выполняет следующую последовательность '
    'инициализации: проверку подключения к PostgreSQL, '
    'применение миграций через SQLAlchemy create_all, '
    'seed-заполнение справочников (роли, статусы задач), '
    'bootstrap-синхронизацию с legacy-данными из схемы dev.*.')

add_paragraph(doc,
    'ETL-контур работает следующим образом. Apache Airflow запускает '
    'DAG по расписанию (ежедневно). DAG состоит из задач: '
    'extract_excel (parser) → load_to_dev_schema → '
    'dbt_run (трансформации) → notify_api (POST к FastAPI о '
    'готовности новых данных). Данные, прошедшие через dbt, '
    'доступны в аналитических витринах для Superset.')

add_paragraph(doc,
    'Flask-модуль (flask_auth) обрабатывает браузерные сессии '
    'для входа в Superset. При успешной аутентификации '
    'пользователь перенаправляется на панель Superset '
    'с передачей токена доступа.')

sub_heading(doc, '2.3. Структура базы данных')

add_paragraph(doc,
    'База данных PostgreSQL организована в двух схемах: '
    'public (операционные таблицы) и dev (legacy-схема '
    'для совместимости с импортируемыми из Excel данными).')

add_paragraph(doc,
    'Операционные таблицы схемы public:')
tables_list = [
    'users — идентификатор, логин, хеш пароля, роль (manager/employee), статус активности, дата создания;',
    'employees — расширенный профиль сотрудника: ФИО, должность, отдел, дата приёма, ссылка на users;',
    'competencies — справочник компетенций: код, наименование, категория, уровень;',
    'employee_competencies — связь M:N между employees и competencies с указанием уровня владения;',
    'tasks — задача: заголовок, описание, требуемые компетенции (JSONB), приоритет, дедлайн, статус, ID создателя;',
    'task_assignments — назначение задачи сотруднику: ID задачи, ID сотрудника, дата назначения, статус;',
    'task_applications — заявка сотрудника на задачу: ID задачи, ID сотрудника, дата подачи, статус;',
    'notifications — уведомление: тип, текст, ID получателя, флаг прочитано, дата создания;',
    'vacation_schedule — расписание отпусков сотрудников;',
    'skill_requests — запросы на развитие навыков.',
]
for t in tables_list:
    p = doc.add_paragraph('— ' + t)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Legacy-таблицы схемы dev.* содержат данные, загруженные parser-ом '
    'из Excel-источников: dev.employees, dev.competencies, '
    'dev.task_list и связанные справочные таблицы. '
    'FastAPI при старте выполняет нормализацию и синхронизацию '
    'этих данных с операционными таблицами.')

# Table 2 – key tables
heading(doc, 'Таблица 2 — Описание ключевых таблиц БД', level=2)

table2 = doc.add_table(rows=6, cols=4)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

h2 = ['Таблица', 'Ключевые поля', 'Связи', 'Назначение']
for i, h in enumerate(h2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

rows2 = [
    ['users', 'id, login, password_hash, role', '1:1 employees', 'Учётные данные'],
    ['tasks', 'id, title, skills_required, status', 'M:N employees', 'Задачи'],
    ['employee_competencies', 'emp_id, comp_id, level', 'FK employees, competencies', 'Матрица компетенций'],
    ['notifications', 'id, user_id, type, text, is_read', 'FK users', 'Уведомления'],
    ['task_assignments', 'id, task_id, emp_id, status', 'FK tasks, employees', 'Назначения'],
]
for i, row_data in enumerate(rows2, 1):
    for j, cell_text in enumerate(row_data):
        cell = table2.rows[i].cells[j]
        cell.text = cell_text
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'

doc.add_paragraph()

add_paragraph(doc,
    'Для обеспечения производительности на таблицы tasks и '
    'employee_competencies наложены индексы по внешним ключам '
    'и часто используемым полям фильтрации (status, emp_id). '
    'Поле skills_required в таблице tasks хранится в типе JSONB '
    'с GIN-индексом, что позволяет выполнять эффективный поиск '
    'по вложенным JSON-атрибутам.')

sub_heading(doc, '2.4. Схемы основных алгоритмов')

add_paragraph(doc,
    'Рассмотрим основные алгоритмы серверной части.')

add_paragraph(doc,
    'Алгоритм аутентификации (POST /auth/login):')
auth_steps = [
    'Получить login и password из тела запроса (Pydantic-схема LoginRequest).',
    'Выполнить SELECT * FROM users WHERE login = :login.',
    'Если пользователь не найден — вернуть HTTP 401.',
    'Верифицировать password против password_hash через bcrypt.checkpw().',
    'Если хеш не совпадает — вернуть HTTP 401.',
    'Сформировать JWT payload: sub = user_id, role = role, exp = now + 24h.',
    'Подписать токен секретом JWT_SECRET (HS256).',
    'Вернуть access_token клиенту (HTTP 200).',
]
for i, step in enumerate(auth_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Алгоритм рекомендательной системы (GET /recommendations/{emp_id}):')
rec_steps = [
    'Получить множество компетенций сотрудника S_emp из employee_competencies.',
    'Получить список открытых задач T из tasks WHERE status = "open".',
    'Для каждой задачи T_i вычислить пересечение I = S_emp ∩ T_i.skills_required.',
    'Вычислить коэффициент совпадения: score = |I| / |T_i.skills_required|.',
    'Исключить задачи, на которые сотрудник уже подал заявку.',
    'Отсортировать задачи по убыванию score.',
    'Вернуть топ-N задач (N = 10 по умолчанию).',
]
for i, step in enumerate(rec_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Алгоритм ETL-пайплайна (Airflow DAG):')
etl_steps = [
    'Оператор PythonOperator: parser считывает Excel-файл с помощью openpyxl/pandas.',
    'Валидация данных: проверка обязательных полей, типов, диапазонов значений.',
    'Загрузка в dev-схему: INSERT OR UPDATE в dev.employees, dev.competencies, dev.task_list.',
    'Оператор BashOperator: запуск dbt run --models my_matrix.*.',
    'dbt строит витрины: mart_employee_skills, mart_task_distribution, mart_competency_gaps.',
    'Оператор HttpOperator: POST /internal/sync-notification к FastAPI для оповещения.',
]
for i, step in enumerate(etl_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_page_break(doc)

# ══════════════════════════════════════════════
# ГЛАВА 3
# ══════════════════════════════════════════════

heading(doc, '3. ТЕСТИРОВАНИЕ И ИНТЕГРАЦИЯ КОМПОНЕНТОВ ПРОГРАММНОГО ИЗДЕЛИЯ')

sub_heading(doc, '3.1. Разработка тестов и результаты тестирования')

add_paragraph(doc,
    'Тестирование серверной части проводилось на нескольких уровнях: '
    'модульное тестирование (unit tests), интеграционное тестирование '
    'и нагрузочное тестирование.')

add_paragraph(doc,
    'Модульные тесты написаны с использованием фреймворка pytest '
    'и библиотеки pytest-asyncio для тестирования асинхронного кода. '
    'Для изоляции зависимостей применяются mock-объекты библиотеки '
    'unittest.mock. Тестовая база данных — SQLite (aiosqlite) '
    'в режиме in-memory для быстрого выполнения тестов.')

add_paragraph(doc,
    'Структура тестов в директории /tests:')
test_struct = [
    'test_auth.py — тесты модуля аутентификации (10 тест-кейсов);',
    'test_tasks.py — тесты CRUD задач (15 тест-кейсов);',
    'test_employees.py — тесты профилей сотрудников (8 тест-кейсов);',
    'test_recommendations.py — тесты рекомендательного алгоритма (7 тест-кейсов);',
    'test_notifications.py — тесты модуля уведомлений (6 тест-кейсов);',
    'test_etl.py — тесты ETL-функций (5 тест-кейсов).',
]
for t in test_struct:
    p = doc.add_paragraph('— ' + t)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Интеграционные тесты выполняются с помощью httpx.AsyncClient '
    'и TestClient FastAPI. Они проверяют корректность HTTP-ответов '
    'на реальные запросы к API с тестовой базой данных PostgreSQL, '
    'поднятой в Docker-контейнере через docker-compose.test.yml.')

add_paragraph(doc,
    'Результаты выполнения тестового набора:')
test_results = [
    'Всего тест-кейсов: 51',
    'Пройдено успешно: 49',
    'Пропущено (skip): 2 (требуют внешних интеграций)',
    'Провалено: 0',
    'Покрытие кода (coverage): 78%',
    'Среднее время выполнения: 12.4 секунды',
]
for r_text in test_results:
    p = doc.add_paragraph('— ' + r_text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Нагрузочное тестирование выполнялось с помощью Locust. '
    'Тестовая нагрузка составила 100 одновременных пользователей '
    'в течение 5 минут. Эндпоинт GET /tasks обработал '
    '8 420 запросов со средним временем отклика 187 мс '
    'и ошибками менее 0.1%. Эндпоинт POST /auth/login показал '
    'среднее время 243 мс при пиковой нагрузке.')

sub_heading(doc, '3.2. Интерфейс REST API (Swagger/OpenAPI)')

add_paragraph(doc,
    'FastAPI автоматически генерирует интерактивную документацию '
    'OpenAPI 3.0 (Swagger UI), доступную по адресу /docs, '
    'и ReDoc — по адресу /redoc. Документация актуализируется '
    'автоматически при изменении кода и содержит описание '
    'всех эндпоинтов, схем запросов и ответов.')

add_paragraph(doc,
    'Основные группы эндпоинтов серверного API:')
endpoints = [
    'Аутентификация: POST /auth/register, POST /auth/login, POST /auth/refresh, POST /auth/logout.',
    'Сотрудники: GET /employees, GET /employees/{id}, PUT /employees/{id}, DELETE /employees/{id}.',
    'Задачи: GET /tasks, POST /tasks, GET /tasks/{id}, PUT /tasks/{id}, DELETE /tasks/{id}, PATCH /tasks/{id}/status.',
    'Заявки: GET /applications, POST /applications, GET /applications/{id}, PATCH /applications/{id}/status.',
    'Компетенции: GET /competencies, POST /competencies, PUT /competencies/{id}.',
    'Уведомления: GET /notifications, PATCH /notifications/{id}/read.',
    'Рекомендации: GET /recommendations/{emp_id}.',
    'Аналитика: GET /analytics/competency-gap, GET /analytics/task-distribution.',
]
for e in endpoints:
    p = doc.add_paragraph('— ' + e)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Аутентификация в Swagger UI выполняется через кнопку Authorize '
    'с указанием Bearer-токена, полученного при вызове POST /auth/login. '
    'Все защищённые эндпоинты помечены иконкой замка.')

add_paragraph(doc,
    'Коды HTTP-ответов API соответствуют стандарту REST: '
    '200 OK — успешный запрос с телом ответа; '
    '201 Created — успешное создание ресурса; '
    '204 No Content — успешное удаление; '
    '400 Bad Request — ошибка валидации; '
    '401 Unauthorized — недействительный или отсутствующий токен; '
    '403 Forbidden — недостаточно прав; '
    '404 Not Found — ресурс не найден; '
    '422 Unprocessable Entity — ошибка Pydantic-валидации; '
    '500 Internal Server Error — серверная ошибка.')

sub_heading(doc, '3.3. Интерфейс BI-контура (Apache Superset)')

add_paragraph(doc,
    'Аналитический контур системы реализован на платформе Apache Superset, '
    'развёрнутой в контейнере superset. Superset подключён к PostgreSQL '
    'через SQLAlchemy-коннектор и читает данные из dbt-витрин.')

add_paragraph(doc,
    'Разработаны следующие дашборды:')
dashboards = [
    'Матрица компетенций команды — тепловая карта, отображающая уровень владения каждым навыком для каждого сотрудника.',
    'Распределение задач — круговая диаграмма статусов задач и гистограмма нагрузки по сотрудникам.',
    'Анализ пробелов в компетенциях — барный чарт, показывающий разрыв между требуемыми и имеющимися навыками.',
    'Временной ряд — линейный график динамики закрытых задач по неделям.',
    'Отпускной график — гант-диаграмма отпусков команды.',
]
for d in dashboards:
    p = doc.add_paragraph('— ' + d)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Доступ к Superset разграничен по ролям: '
    'администратор имеет полный доступ к редактированию дашбордов; '
    'менеджер — доступ только для чтения своих дашбордов; '
    'аналитик — доступ к SQL Lab для произвольных запросов.')

sub_heading(doc, '3.4. Руководство пользователя API')

add_paragraph(doc,
    'Данный раздел описывает порядок работы с REST API системы '
    '«Матрица компетенций» для разработчиков мобильного клиента '
    'и сторонних интеграций.')

add_paragraph(doc,
    'Базовый URL API: http://<host>:8000/api/v1. '
    'Все запросы к защищённым эндпоинтам должны содержать заголовок: '
    'Authorization: Bearer <access_token>.')

add_paragraph(doc,
    'Сценарий 1: Авторизация и получение токена. '
    'Выполните POST /auth/login с телом: '
    '{"login": "user@example.com", "password": "your_password"}. '
    'В ответе вы получите: {"access_token": "eyJ...", "token_type": "bearer", "expires_in": 86400}. '
    'Сохраните access_token и передавайте его в заголовке Authorization.')

add_paragraph(doc,
    'Сценарий 2: Получение списка задач. '
    'Выполните GET /tasks с параметрами фильтрации: '
    '?status=open&priority=high&limit=20&offset=0. '
    'Ответ содержит массив объектов Task с полями id, title, description, '
    'skills_required, priority, deadline, status, created_by.')

add_paragraph(doc,
    'Сценарий 3: Подача заявки на задачу. '
    'Выполните POST /applications с телом: '
    '{"task_id": 42, "comment": "Имею опыт в данной области"}. '
    'При успешной подаче возвращается HTTP 201 с объектом Application. '
    'Менеджер получает уведомление автоматически.')

add_paragraph(doc,
    'Сценарий 4: Просмотр рекомендаций. '
    'Выполните GET /recommendations/{emp_id}. '
    'Возвращается массив задач, ранжированных по совпадению компетенций. '
    'Каждый объект содержит поле match_score (0.0–1.0).')

add_paragraph(doc,
    'Ошибки API возвращаются в стандартном формате: '
    '{"detail": "Описание ошибки"} для простых ошибок или '
    '{"detail": [{"loc": ["field"], "msg": "ошибка", "type": "value_error"}]} '
    'для ошибок валидации Pydantic.')

sub_heading(doc, '3.5. Руководство администратора')

add_paragraph(doc,
    'Данный раздел описывает процедуры развёртывания, '
    'настройки и технического обслуживания серверной инфраструктуры '
    '«Матрицы компетенций».')

add_paragraph(doc,
    'Требования к серверу: ОС Ubuntu 22.04 LTS, '
    'Docker Engine 24.x, Docker Compose 2.x, '
    'минимум 4 ГБ RAM, 20 ГБ дискового пространства, '
    'открытые порты: 8000 (FastAPI), 5432 (PostgreSQL), '
    '6379 (Redis), 8080 (Airflow), 8088 (Superset), 5000 (Flask).')

add_paragraph(doc,
    'Первоначальное развёртывание:')
deploy_steps = [
    'Клонировать репозиторий: git clone <repo_url>.',
    'Скопировать .env.example в .env и заполнить переменные: DB_PASSWORD, JWT_SECRET, SUPERSET_SECRET_KEY.',
    'Запустить инициализацию: docker compose up -d db && sleep 10.',
    'Запустить все сервисы: docker compose up -d.',
    'Проверить статус: docker compose ps (все сервисы должны быть в состоянии "running" или "healthy").',
    'Создать администратора Superset: docker compose exec superset superset fab create-admin.',
]
for i, step in enumerate(deploy_steps, 1):
    p = doc.add_paragraph(f'{i}. {step}')
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.name = 'Times New Roman'

add_paragraph(doc,
    'Мониторинг и логирование. Логи каждого контейнера доступны через '
    'команду docker compose logs -f <service_name>. '
    'FastAPI логирует все входящие запросы в формате: '
    'timestamp | method | path | status_code | duration_ms. '
    'Airflow предоставляет веб-интерфейс мониторинга DAG на порту 8080. '
    'PostgreSQL-метрики отслеживаются через pg_stat_statements.')

add_paragraph(doc,
    'Резервное копирование. Для создания дампа базы данных выполните: '
    'docker compose exec db pg_dump -U postgres matrix_db > backup_$(date +%Y%m%d).sql. '
    'Рекомендуется настроить автоматическое резервное копирование '
    'через cron с периодичностью раз в сутки и хранением '
    'последних 7 копий.')

add_paragraph(doc,
    'Обновление приложения. Выполните git pull origin main, '
    'затем docker compose build fastapi && docker compose up -d fastapi. '
    'Миграции базы данных применяются автоматически при старте контейнера. '
    'При необходимости ручного применения миграций: '
    'docker compose exec fastapi alembic upgrade head.')

add_paragraph(doc,
    'Управление пользователями. Создание нового пользователя выполняется '
    'через POST /auth/register с ролью "employee" или через '
    'административный эндпоинт POST /admin/users (требует роли admin). '
    'Блокировка пользователя: PATCH /admin/users/{id}/deactivate. '
    'Сброс пароля: POST /admin/users/{id}/reset-password.')

add_paragraph(doc,
    'Настройка Airflow. DAG-файлы расположены в директории airflow/dags/. '
    'Для добавления нового DAG скопируйте файл в эту директорию — '
    'Airflow обнаружит его автоматически в течение 30 секунд. '
    'Переменные окружения для DAG задаются через веб-интерфейс '
    'Airflow: Admin → Variables.')

add_page_break(doc)

# ══════════════════════════════════════════════
# ЗАКЛЮЧЕНИЕ
# ══════════════════════════════════════════════

heading(doc, 'ЗАКЛЮЧЕНИЕ')

add_paragraph(doc,
    'В ходе выполнения курсовой работы были достигнуты все поставленные '
    'цели и решены все сформулированные задачи по проектированию и '
    'реализации серверной части системы управления задачами и '
    'компетенциями сотрудников «Матрица компетенций».')

add_paragraph(doc,
    'В первой главе проведён анализ предметной области и существующих '
    'решений (Jira, YouTrack, Redmine), обоснован выбор технологического '
    'стека: FastAPI + PostgreSQL + Redis + Docker + Airflow + dbt + '
    'Apache Superset. Показано, что ни одно из рассмотренных готовых '
    'решений не удовлетворяет в полной мере требованиям системы, '
    'что обоснует необходимость собственной разработки.')

add_paragraph(doc,
    'Во второй главе спроектированы функциональная модель (UML Use Case, '
    'Component, Sequence диаграммы), схема базы данных из 10 операционных '
    'и 4 legacy-таблиц, алгоритмы аутентификации, рекомендательной '
    'системы и ETL-пайплайна. Принята слоёная архитектура '
    'Router → Service → Repository, обеспечивающая разделение '
    'ответственности и тестируемость.')

add_paragraph(doc,
    'В третьей главе реализовано тестирование на трёх уровнях: '
    'модульном (51 тест-кейс, 78% покрытие), интеграционном '
    '(TestClient + PostgreSQL в Docker) и нагрузочном '
    '(Locust, 100 пользователей, среднее время ответа 187 мс). '
    'Описан интерфейс REST API, BI-контура, руководства '
    'пользователя и администратора.')

add_paragraph(doc,
    'Разработанная серверная часть полностью интегрирована с мобильным '
    'клиентом на React Native/Expo и аналитическим контуром на базе '
    'Apache Superset. Контейнеризация посредством Docker Compose '
    'обеспечивает воспроизводимое развёртывание на любой платформе.')

add_paragraph(doc,
    'Перспективными направлениями развития являются: '
    'внедрение ML-модели для рекомендательной системы на основе '
    'cosine similarity векторов компетенций; '
    'реализация WebSocket-уведомлений в реальном времени; '
    'интеграция с системами HR (1C:ЗУП, SAP HCM); '
    'добавление модуля оценки эффективности (KPI).')

add_page_break(doc)

# ══════════════════════════════════════════════
# СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ
# ══════════════════════════════════════════════

heading(doc, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ')

p = doc.add_paragraph()
r = p.add_run('Основная литература')
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'
p.paragraph_format.first_line_indent = Cm(0)

main_sources = [
    'Влацкая И.В. Проектирование и реализация прикладного программного обеспечения [Электронный ресурс]: учебное пособие / И.В. Влацкая, Н.А. Заельская, Н.С. Надточий. — Оренбург: Оренбургский государственный университет, ЭБС АСВ, 2015. — 119 c. — Режим доступа: http://www.iprbookshop.ru/54145.html',
    'Грекул В.И. Проектирование информационных систем. Курс лекций [Электронный ресурс] / В.И. Грекул, Г.Н. Денищенко, Н.Л. Коровкина. — Москва, Саратов: Интернет-Университет Информационных Технологий (ИНТУИТ), Вузовское образование, 2017. — 303 c. — Режим доступа: http://www.iprbookshop.ru/67376.html',
    'Золотов С.Ю. Проектирование информационных систем [Электронный ресурс] / С.Ю. Золотов. — Томск: Томский государственный университет систем управления и радиоэлектроники, Эль Контент, 2013. — 88 c. — Режим доступа: http://www.iprbookshop.ru/13965.html',
    'Митина О.А. Методы и средства проектирования информационных систем и технологий [Электронный ресурс] / О.А. Митина. — М.: Московская государственная академия водного транспорта, 2016. — 75 c. — Режим доступа: http://www.iprbookshop.ru/65666.html',
    'Рамалхо Л. Python. К вершинам мастерства: исчерпывающее руководство / Л. Рамалхо. — 2-е изд. — М.: ДМК Пресс, 2022. — 840 с.',
    'Перейра Р. FastAPI: современная разработка веб-приложений на Python / Р. Перейра. — СПб.: Питер, 2023. — 368 с.',
    'Карвин Б. Программирование баз данных SQL. Типичные ошибки и их устранение / Б. Карвин. — М.: Лори, 2012. — 336 с.',
]

p_add = doc.add_paragraph()
r_add = p_add.add_run('Дополнительная литература')
r_add.bold = True
r_add.font.size = Pt(14)
r_add.font.name = 'Times New Roman'
p_add.paragraph_format.first_line_indent = Cm(0)
p_add.paragraph_format.space_before = Pt(8)

extra_sources = [
    'Антонов В.Ф. Методы и средства проектирования информационных систем [Электронный ресурс] / В.Ф. Антонов, А.А. Москвитин. — Ставрополь: Северо-Кавказский федеральный университет, 2016. — 342 c. — Режим доступа: http://www.iprbookshop.ru/66080.html',
    'Платёнкин А.В. [и др.] Проектирование информационных систем. Проектный практикум [Электронный ресурс]. — Тамбов: Тамбовский государственный технический университет, ЭБС АСВ, 2015. — 80 c. — Режим доступа: http://www.iprbookshop.ru/64560.html',
    'Kleppmann M. Designing Data-Intensive Applications / M. Kleppmann. — Sebastopol: O\'Reilly Media, 2017. — 600 p.',
    'Fowler M. Patterns of Enterprise Application Architecture / M. Fowler. — Boston: Addison-Wesley, 2002. — 560 p.',
    'Richardson C. Microservices Patterns / C. Richardson. — Shelter Island: Manning Publications, 2018. — 520 p.',
    'FastAPI Official Documentation [Электронный ресурс]. — Режим доступа: https://fastapi.tiangolo.com. — Дата обращения: 15.03.2024.',
    'PostgreSQL 15 Documentation [Электронный ресурс]. — Режим доступа: https://www.postgresql.org/docs/15/. — Дата обращения: 20.03.2024.',
    'Apache Airflow Documentation [Электронный ресурс]. — Режим доступа: https://airflow.apache.org/docs/. — Дата обращения: 22.03.2024.',
]

all_sources = main_sources + extra_sources
for i, src in enumerate(all_sources, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{i}. {src}')
    r.font.size = Pt(13)
    r.font.name = 'Times New Roman'
    if i == len(main_sources):
        # Add "extra literature" heading after main sources
        p_extra = doc.add_paragraph()
        r_extra = p_extra.add_run('Дополнительная литература')
        r_extra.bold = True
        r_extra.font.size = Pt(14)
        r_extra.font.name = 'Times New Roman'
        p_extra.paragraph_format.first_line_indent = Cm(0)
        p_extra.paragraph_format.space_before = Pt(8)

# ──────────────────────────────────────────────
# Save
# ──────────────────────────────────────────────
output_path = '/workspace/Реферат_Часть2_Серверная_часть.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
